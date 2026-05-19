"""Render the LTCMA markdown report into Word (.docx) and PDF.
Word  : python-docx (custom markdown parser)
PDF   : markdown -> styled HTML -> weasyprint
"""
import os, re
import markdown as md
from weasyprint import HTML
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

REP = os.path.expanduser("~/LTCMA/report")
SRC = f"{REP}/LTCMA_2026.md"
NAVY = RGBColor(0x1A, 0x3A, 0x5C)
text = open(SRC, encoding="utf-8").read()
lines = text.split("\n")

# ============================================================ WORD
doc = Document()
st = doc.styles["Normal"].font
st.name, st.size = "Calibri", Pt(10.5)

INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")

def add_runs(par, s):
    for tok in INLINE.split(s):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            par.add_run(tok[2:-2]).bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = par.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9)
        elif tok.startswith("*") and tok.endswith("*"):
            par.add_run(tok[1:-1]).italic = True
        else:
            par.add_run(tok)

def flush_table(rows):
    if not rows:
        return
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    cells = [c for c in cells if not all(set(x) <= set("-: ") for x in c)]
    if not cells:
        return
    t = doc.add_table(rows=len(cells), cols=len(cells[0]))
    t.style = "Light Grid Accent 1"
    for i, row in enumerate(cells):
        for j, val in enumerate(row):
            cell = t.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs(p, val.replace("**", ""))
            for run in p.runs:
                run.font.size = Pt(8.5)
                if i == 0:
                    run.bold = True
    doc.add_paragraph()

i, tbl, code, codebuf = 0, [], False, []
while i < len(lines):
    ln = lines[i]
    if ln.strip().startswith("```"):
        if code:
            p = doc.add_paragraph()
            r = p.add_run("\n".join(codebuf)); r.font.name = "Consolas"; r.font.size = Pt(9)
            codebuf, code = [], False
        else:
            code = True
        i += 1; continue
    if code:
        codebuf.append(ln); i += 1; continue
    if ln.strip().startswith("|"):
        tbl.append(ln); i += 1; continue
    if tbl:
        flush_table(tbl); tbl = []
    if ln.startswith("# "):
        h = doc.add_heading(level=0); add_runs(h, ln[2:])
    elif ln.startswith("## "):
        h = doc.add_heading(level=1); add_runs(h, ln[3:])
        for r in h.runs: r.font.color.rgb = NAVY
    elif ln.startswith("### "):
        h = doc.add_heading(level=2); add_runs(h, ln[4:])
        for r in h.runs: r.font.color.rgb = NAVY
    elif ln.startswith("> "):
        p = doc.add_paragraph(style="Intense Quote"); add_runs(p, ln[2:])
    elif re.match(r"^!\[.*\]\(.*\)", ln):
        m = re.match(r"^!\[(.*)\]\((.*)\)", ln)
        img = os.path.join(REP, m.group(2))
        if os.path.exists(img):
            doc.add_picture(img, width=Inches(6.3))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph(m.group(1))
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in cap.runs:
                r.italic = True; r.font.size = Pt(8)
    elif re.match(r"^[-*] ", ln):
        p = doc.add_paragraph(style="List Bullet"); add_runs(p, ln[2:])
    elif re.match(r"^\d+\. ", ln):
        p = doc.add_paragraph(style="List Number"); add_runs(p, ln.split(". ", 1)[1])
    elif ln.strip() == "---":
        pass
    elif ln.strip():
        p = doc.add_paragraph(); add_runs(p, ln)
    i += 1
if tbl:
    flush_table(tbl)
doc.save(f"{REP}/LTCMA_2026.docx")
print("wrote LTCMA_2026.docx")

# ============================================================ PDF
html_body = md.markdown(text, extensions=["tables", "fenced_code", "sane_lists"])
CSS = """
@page { size: Letter; margin: 1.8cm 2cm;
  @bottom-center { content: counter(page); font-size: 8pt; color: #888; } }
body { font-family: 'Helvetica','Arial',sans-serif; font-size: 10pt;
  line-height: 1.45; color: #1f1f1f; }
h1 { font-size: 21pt; color: #1a3a5c; margin: 0 0 4pt; }
h1 + p { color: #555; }
h2 { font-size: 14pt; color: #1a3a5c; border-bottom: 2px solid #1a3a5c;
  padding-bottom: 3pt; margin-top: 22pt; }
h3 { font-size: 11.5pt; color: #1a3a5c; margin-top: 14pt; }
table { border-collapse: collapse; width: 100%; margin: 8pt 0; font-size: 8.4pt; }
th { background: #1a3a5c; color: #fff; padding: 5pt 6pt; text-align: left; }
td { border: 1px solid #d0d5da; padding: 4pt 6pt; }
tr:nth-child(even) td { background: #f4f6f8; }
blockquote { border-left: 3px solid #1a3a5c; background: #eef2f6;
  margin: 8pt 0; padding: 6pt 12pt; font-style: italic; color: #444; }
code { background: #eef0f2; font-family: 'Consolas',monospace; font-size: 8.6pt;
  padding: 1pt 3pt; }
pre { background: #f4f5f6; border-left: 3px solid #b0b8c0; padding: 8pt;
  font-size: 8.6pt; white-space: pre-wrap; }
hr { border: none; border-top: 1px solid #d0d5da; margin: 14pt 0; }
strong { color: #11243a; }
"""
CSS += "\nimg { width: 100%; margin: 6pt 0; }\n"
HTML(string=f"<style>{CSS}</style>{html_body}",
     base_url=REP + "/").write_pdf(f"{REP}/LTCMA_2026.pdf")
print("wrote LTCMA_2026.pdf")
